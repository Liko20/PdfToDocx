import os
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from pdf2image import convert_from_path
import pytesseract
from docx import Document

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads/'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Specify the path to Tesseract executable if necessary
pytesseract.pytesseract.tesseract_cmd = r'C:/Program Files/Tesseract-OCR/tesseract.exe'

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'pdf' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['pdf']

    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    # Normalize the file extension to be case insensitive
    file_extension = file.filename.lower().endswith('.pdf')
    
    if file and file_extension:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        docx_path, analysis = convert_pdf_to_docx(file_path)
        return jsonify({
            'message': 'File uploaded and converted successfully!',
            'docx_path': docx_path,
            'analysis': analysis
        }), 200
    else:
        return jsonify({'error': 'Invalid file format'}), 400

def convert_pdf_to_docx(pdf_path):
    images = convert_from_path(pdf_path)
    doc = Document()
    
    # Extract text using pytesseract
    text = ""
    for image in images:
        text += pytesseract.image_to_string(image)
    
    # Create a new DOCX file
    for line in text.split('\n'):
        if line.strip() != '':
            doc.add_paragraph(line)
    
    # Save DOCX file
    docx_path = pdf_path.replace('.pdf', '.docx').replace('.PDF', '.docx')
    doc.save(docx_path)
    
    # Analyze the DOCX file
    analysis = analyze_docx(doc)
    
    return docx_path, analysis

def analyze_docx(doc):
    text = '\n'.join([p.text for p in doc.paragraphs])
    num_sentences = text.count('.') + text.count('!') + text.count('?')
    num_words = len(text.split())
    num_characters = len(text)
    num_lines = text.count('\n') + 1
    
    return {
        'num_sentences': num_sentences,
        'num_words': num_words,
        'num_characters': num_characters,
        'num_lines': num_lines
    }

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    return send_file(os.path.join(app.config['UPLOAD_FOLDER'], filename))

if __name__ == '__main__':
    app.run(debug=True)
